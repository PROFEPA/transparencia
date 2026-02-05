"""
Extractor de datos para archivos Word (DOCX).
Procesa documentos narrativos para extraer definiciones y metodologías.
"""
from docx import Document
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import re
import logging
from slugify import slugify

from models import Indicator, DataQualityReport

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extrae información de documentos Word."""
    
    def __init__(self, file_path: Path, source_config: dict):
        self.file_path = file_path
        self.config = source_config
        self.quality_report = DataQualityReport(
            archivo_fuente=str(file_path),
            fecha_extraccion=datetime.now().isoformat()
        )
        self.indicators: List[Indicator] = []
        self.narrative_content: Dict[str, str] = {}
        self.sections: Dict[str, str] = {}
    
    def extract(self) -> Tuple[List[Indicator], Dict[str, str], DataQualityReport]:
        """Extrae indicadores y contenido narrativo del documento."""
        try:
            logger.info(f"Procesando documento: {self.file_path}")
            
            doc = Document(self.file_path)
            
            # Extraer todo el texto estructurado
            full_text = self._extract_full_text(doc)
            
            # Detectar secciones
            self._detect_sections(full_text)
            
            # Extraer indicadores de las secciones
            self._extract_indicators_from_sections()
            
            # Extraer tablas si existen
            self._extract_from_tables(doc)
            
            logger.info(f"Extracción completada: {len(self.indicators)} indicadores, {len(self.sections)} secciones")
            
        except Exception as e:
            error_msg = f"Error procesando {self.file_path}: {str(e)}"
            logger.error(error_msg)
            self.quality_report.add_error(error_msg)
        
        return self.indicators, self.sections, self.quality_report
    
    def _extract_full_text(self, doc: Document) -> str:
        """Extrae todo el texto del documento preservando estructura."""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detectar títulos por estilo
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name or "título" in style_name:
                    text = f"\n### {text} ###\n"
                paragraphs.append(text)
        
        return "\n".join(paragraphs)
    
    def _detect_sections(self, text: str):
        """Detecta y extrae secciones del documento."""
        # Patrones para detectar secciones relevantes
        section_patterns = [
            (r"(?:objetivo|propósito|proposito)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "objetivo"),
            (r"(?:fin|finalidad)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "fin"),
            (r"(?:indicador(?:es)?|medición)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "indicadores"),
            (r"(?:definición|definicion|concepto)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "definicion"),
            (r"(?:método de cálculo|metodo de calculo|fórmula|formula)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "metodo_calculo"),
            (r"(?:unidad de medida|unidad)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "unidad_medida"),
            (r"(?:meta|metas)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "meta"),
            (r"(?:resumen narrativo)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "resumen_narrativo"),
            (r"(?:programa presupuestario|clave presupuestaria)[\s:]*(.+?)(?=\n###|\n\n\n|$)", "programa"),
            (r"(?:componente)[\s:]*\d*[\s:]*(.+?)(?=\n###|\n\n\n|$)", "componente"),
            (r"(?:actividad)[\s:]*\d*[\s:]*(.+?)(?=\n###|\n\n\n|$)", "actividad"),
        ]
        
        for pattern, section_name in section_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                # Limpiar y guardar
                content = " ".join(matches).strip()
                content = re.sub(r'\s+', ' ', content)
                content = content[:2000]  # Limitar longitud
                self.sections[section_name] = content
        
        # Guardar texto completo como narrativa general
        self.narrative_content["texto_completo"] = text[:10000]
    
    def _extract_indicators_from_sections(self):
        """Extrae indicadores de las secciones detectadas."""
        # Buscar patrones de indicadores en el texto
        indicator_pattern = r"(?:indicador|medición)[\s:]+([^.]+\.)"
        
        text = self.narrative_content.get("texto_completo", "")
        matches = re.findall(indicator_pattern, text, re.IGNORECASE)
        
        for i, match in enumerate(matches):
            nombre = match.strip()
            if len(nombre) < 10 or len(nombre) > 500:
                continue
            
            indicator_id = self._generate_indicator_id(nombre)
            
            # Verificar duplicados
            if any(ind.id == indicator_id for ind in self.indicators):
                continue
            
            indicator = Indicator(
                id=indicator_id,
                nombre=nombre,
                programa=self.config.get("programa", "G014"),
                anio=self.config.get("anio", 2026),
                fuente=f"{self.file_path.name} - Sección narrativa",
                definicion=self.sections.get("definicion"),
                metodo_calculo=self.sections.get("metodo_calculo"),
                unidad_medida=self.sections.get("unidad_medida"),
                nivel=self._detect_nivel_from_context(nombre),
                notas="Extraído de documento narrativo",
                ultima_actualizacion=datetime.now().strftime("%Y-%m-%d")
            )
            
            self.indicators.append(indicator)
            self.quality_report.filas_validas += 1
    
    def _extract_from_tables(self, doc: Document):
        """Extrae indicadores de tablas en el documento."""
        for table_idx, table in enumerate(doc.tables):
            try:
                headers = []
                
                # Primera fila como encabezados
                if table.rows:
                    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                
                # Buscar columna de indicadores
                indicator_col_idx = None
                for i, header in enumerate(headers):
                    if any(kw in header for kw in ["indicador", "nombre", "descripción", "descripcion"]):
                        indicator_col_idx = i
                        break
                
                if indicator_col_idx is None:
                    continue
                
                # Procesar filas de datos
                for row_idx, row in enumerate(table.rows[1:], start=2):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if len(cells) > indicator_col_idx:
                        nombre = cells[indicator_col_idx]
                        
                        if not nombre or len(nombre) < 5:
                            continue
                        
                        indicator_id = self._generate_indicator_id(nombre)
                        
                        # Verificar duplicados
                        if any(ind.id == indicator_id for ind in self.indicators):
                            continue
                        
                        # Extraer otros campos si existen
                        definicion = None
                        metodo = None
                        unidad = None
                        
                        for i, header in enumerate(headers):
                            if i < len(cells):
                                if "definición" in header or "definicion" in header:
                                    definicion = cells[i]
                                elif "método" in header or "metodo" in header or "fórmula" in header:
                                    metodo = cells[i]
                                elif "unidad" in header:
                                    unidad = cells[i]
                        
                        indicator = Indicator(
                            id=indicator_id,
                            nombre=nombre,
                            programa=self.config.get("programa", "G014"),
                            anio=self.config.get("anio", 2026),
                            fuente=f"{self.file_path.name} - Tabla {table_idx + 1}, Fila {row_idx}",
                            definicion=definicion or self.sections.get("definicion"),
                            metodo_calculo=metodo or self.sections.get("metodo_calculo"),
                            unidad_medida=unidad or self.sections.get("unidad_medida"),
                            nivel=self._detect_nivel_from_context(nombre),
                            notas="Extraído de tabla en documento",
                            ultima_actualizacion=datetime.now().strftime("%Y-%m-%d")
                        )
                        
                        self.indicators.append(indicator)
                        self.quality_report.filas_validas += 1
            
            except Exception as e:
                self.quality_report.add_warning(f"Error en tabla {table_idx}: {str(e)}")
    
    def _generate_indicator_id(self, nombre: str) -> str:
        """Genera un ID único para el indicador."""
        base = slugify(nombre[:50], lowercase=True, separator="-")
        programa = self.config.get("programa", "g014").lower()
        anio = self.config.get("anio", 2026)
        return f"{base}-{programa}-{anio}"
    
    def _detect_nivel_from_context(self, nombre: str) -> Optional[str]:
        """Detecta el nivel MIR basándose en el contexto."""
        nombre_lower = nombre.lower()
        
        if "fin" in nombre_lower[:20]:
            return "Fin"
        elif "propósito" in nombre_lower or "proposito" in nombre_lower:
            return "Propósito"
        elif "componente" in nombre_lower:
            return "Componente"
        elif "actividad" in nombre_lower:
            return "Actividad"
        
        # Usar sección del documento si está disponible
        if "fin" in self.sections:
            return "Fin"
        elif "componente" in self.sections:
            return "Componente"
        elif "actividad" in self.sections:
            return "Actividad"
        
        return None


def extract_from_docx(file_path: Path, config: dict) -> Tuple[List[Indicator], Dict[str, str], DataQualityReport]:
    """Función de conveniencia para extraer de DOCX."""
    extractor = DocxExtractor(file_path, config)
    return extractor.extract()
